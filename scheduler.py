# ================= FORCE IPV4 (CRITICAL FIX) =================
import socket
def force_ipv4():
    socket._orig_getaddrinfo = socket.getaddrinfo
    socket.getaddrinfo = lambda *args, **kwargs: [
        addr for addr in socket._orig_getaddrinfo(*args, **kwargs)
        if addr[0] == socket.AF_INET
    ]
force_ipv4()
# ============================================================
import os
import time
import json
import shutil
import logging
import warnings
from datetime import datetime, time as dt_time
from docx import Document
from telegram import InlineKeyboardButton, InlineKeyboardMarkup, Update
from telegram.ext import (
    ApplicationBuilder,
    CallbackQueryHandler,
    ContextTypes,
    MessageHandler,
    filters,
)
from reports import register_reports
from config import WORD_FILE, CHAT_ID, IST, BOT_TOKEN
from drive import upload_to_drive
from telegram.warnings import PTBUserWarning
warnings.filterwarnings("ignore", category=PTBUserWarning)
warnings.filterwarnings("ignore", category=FutureWarning)
# ================= LOGGING =================
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s"
)
logging.getLogger("httpx").setLevel(logging.WARNING)
logging.getLogger("apscheduler").setLevel(logging.WARNING)
logging.getLogger("telegram").setLevel(logging.WARNING)
# ================= CONSTANTS =================
START_DATE = "2026-01-12"
STATE_FILE = "bot_state.json"
EVENING_RETRY_JOB = "evening_retry"
HARD_TOPIC_TIMEOUT_JOB = "hard_topic_timeout"
REQUIRED_COLUMNS = {"status", "hard topic"}
# ================= INIT WORD FILE =================
word_dir = os.path.dirname(WORD_FILE)
if word_dir and not os.path.exists(word_dir):
    os.makedirs(word_dir, exist_ok=True)  # 🔴 FIX: create /data
if not os.path.exists(WORD_FILE):
    shutil.copy("placepment_plan.docx", WORD_FILE)
    print("📄 Word file copied to volume")
# ================= STATE HANDLING =================
def load_state():
    if not os.path.exists(STATE_FILE):
        return {}
    with open(STATE_FILE, "r") as f:
        return json.load(f)
def save_state(state):
    with open(STATE_FILE, "w") as f:
        json.dump(state, f)
# ================= WORD SAFETY =================
def safe_open_docx(path):
    try:
        return Document(path)
    except Exception:
        logging.exception("Failed to open Word file")
        return None
def safe_save_docx(doc, path):
    try:
        doc.save(path)
        return True
    except Exception:
        logging.exception("Failed to save Word file")
        return False
def validate_word_structure():
    doc = safe_open_docx(WORD_FILE)
    if not doc or not doc.tables:
        raise RuntimeError("Word file has no tables")
    headers = {
        cell.text.strip().lower()
        for cell in doc.tables[0].rows[0].cells
    }
    missing = REQUIRED_COLUMNS - headers
    if missing:
        raise RuntimeError(f"Missing required columns: {missing}")
# ================= HELPERS =================
def upload_with_retry(local_file, filename, retries=3, delay=2):
    for attempt in range(1, retries + 1):
        if upload_to_drive(local_file, filename):
            logging.info("Drive sync successful (attempt %s)", attempt)
            return True
        logging.warning("Drive sync failed (attempt %s/%s)", attempt, retries)
        time.sleep(delay)
    logging.error("Drive sync failed after all retries")
    return False
def get_day_number():
    start = datetime.strptime(START_DATE, "%Y-%m-%d").date()
    today = datetime.now(IST).date()
    if today < start:
        return None
    return (today - start).days + 1
def get_table_and_row(doc, day):
    table_index = (day - 1) // 7
    row_index = ((day - 1) % 7) + 1
    if table_index >= len(doc.tables):
        return None, None
    return doc.tables[table_index], row_index
def find_column_index(table, column_name):
    for i, cell in enumerate(table.rows[0].cells):
        if cell.text.strip().lower() == column_name.lower():
            return i
    return None
# ================= WORD UPDATE =================
def update_status_in_word(symbol):
    day = get_day_number()
    if day is None:
        return
    doc = safe_open_docx(WORD_FILE)
    if not doc:
        return
    table, row = get_table_and_row(doc, day)
    if not table or row >= len(table.rows):
        return
    status_col = find_column_index(table, "status")
    if status_col is None:
        return
    table.rows[row].cells[status_col].text = symbol
    if not safe_save_docx(doc, WORD_FILE):
        return
    upload_with_retry(WORD_FILE, "placepment_plan.docx")
    logging.info("Status updated for day %s", day)
# ================= EVENING =================
async def evening_buttons(context: ContextTypes.DEFAULT_TYPE):
    retries = context.bot_data.get("evening_retry_count", 0)
    if retries >= 2:
        context.bot_data["evening_retry_count"] = 0
        return
    keyboard = [[
        InlineKeyboardButton("✅ Yes", callback_data="evening_yes"),
        InlineKeyboardButton("❌ No", callback_data="evening_no"),
    ]]
    await context.bot.send_message(
        chat_id=CHAT_ID,
        text="Have you started studying 📔 ?",
        reply_markup=InlineKeyboardMarkup(keyboard),
    )
    context.bot_data["evening_retry_count"] = retries + 1
    for job in context.job_queue.get_jobs_by_name(EVENING_RETRY_JOB):
        job.schedule_removal()
    context.job_queue.run_once(
        evening_buttons,
        when=300,  # 5 minutes (FIXED)
        name=EVENING_RETRY_JOB
    )
# ================= NIGHT =================
async def night_buttons(context: ContextTypes.DEFAULT_TYPE):
    keyboard = [[
        InlineKeyboardButton("✅ Yes", callback_data="night_yes"),
        InlineKeyboardButton("❌ No", callback_data="night_no"),
    ]]
    await context.bot.send_message(
        chat_id=CHAT_ID,
        text="Did you complete today’s portion?",
        reply_markup=InlineKeyboardMarkup(keyboard),
    )
# ================= CALLBACK =================
async def button_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    state = load_state()
    if query.data == "evening_yes":
        context.bot_data["evening_retry_count"] = 0
        for job in context.job_queue.get_jobs_by_name(EVENING_RETRY_JOB):
            job.schedule_removal()
        await query.edit_message_text("👍 Good, start studying 💪")
    elif query.data == "evening_no":
        await query.edit_message_text("⏳ Okay, I’ll remind you again in 5 minutes.")
    elif query.data == "night_yes":
        update_status_in_word("✅")
        state["awaiting_hard_topic"] = True
        save_state(state)
        for job in context.job_queue.get_jobs_by_name(HARD_TOPIC_TIMEOUT_JOB):
            job.schedule_removal()
        context.job_queue.run_once(
            hard_topic_timeout,
            when=120,
            name=HARD_TOPIC_TIMEOUT_JOB,
            chat_id=query.message.chat_id,
        )
        await query.edit_message_text(
            "🎉 Marked as COMPLETED ✅\n\nWhich topic did you find hard today?"
        )
    elif query.data == "night_no":
        update_status_in_word("❌")
        await query.edit_message_text(
            "⚠️ Marked as NOT completed ❌\nTry again tomorrow 💪"
        )
# ================= HARD TOPIC =================
async def hard_topic_timeout(context: ContextTypes.DEFAULT_TYPE):
    state = load_state()
    if not state.get("awaiting_hard_topic"):
        return
    day = get_day_number()
    if day is None:
        return
    doc = safe_open_docx(WORD_FILE)
    if not doc:
        return
    table, row = get_table_and_row(doc, day)
    if not table:
        return
    hard_col = find_column_index(table, "hard topic")
    if hard_col is not None and row < len(table.rows):
        table.rows[row].cells[hard_col].text = "None"
        safe_save_docx(doc, WORD_FILE)
        upload_with_retry(WORD_FILE, "placepment_plan.docx")
    state["awaiting_hard_topic"] = False
    save_state(state)
    await context.bot.send_message(
        chat_id=context.job.chat_id,
        text="⏰ No response received. Hard topic marked as None."
    )
async def hard_topic_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    state = load_state()
    if not state.get("awaiting_hard_topic"):
        return
    topic = update.message.text
    day = get_day_number()
    if day is None:
        return
    doc = safe_open_docx(WORD_FILE)
    if not doc:
        return
    table, row = get_table_and_row(doc, day)
    if not table:
        return
    hard_col = find_column_index(table, "hard topic")
    if hard_col is not None and row < len(table.rows):
        table.rows[row].cells[hard_col].text = topic
        safe_save_docx(doc, WORD_FILE)
        upload_with_retry(WORD_FILE, "placepment_plan.docx")
    state["awaiting_hard_topic"] = False
    save_state(state)
    await update.message.reply_text("📝 Hard topic saved successfully ✅")
# ================= APP =================
validate_word_structure()
app = ApplicationBuilder().token(BOT_TOKEN).build()
app.add_handler(CallbackQueryHandler(button_callback))
app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, hard_topic_handler))
register_reports(app)
# ================= SCHEDULE =================
ALL_DAYS = (0, 1, 2, 3, 4, 5, 6)
app.job_queue.run_daily(
    evening_buttons,
    time=dt_time(hour=17, minute=30, tzinfo=IST),
    days=ALL_DAYS
)
app.job_queue.run_daily(
    night_buttons,
    time=dt_time(hour=22, minute=30, tzinfo=IST),
    days=ALL_DAYS
)
# ================= START =================
logging.info("Bot running...")
app.run_polling(drop_pending_updates=True)
