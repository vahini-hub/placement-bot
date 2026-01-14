from datetime import datetime, timedelta, time as dt_time
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from telegram import Update
from telegram.ext import ContextTypes, CommandHandler,JobQueue
from config import WORD_FILE, CHAT_ID, IST
import matplotlib.pyplot as plt
from collections import Counter

def job_wrapper(async_func):
    async def wrapped(context):
        try:
            await async_func(context)
        except Exception as e:
            pass
    return wrapped
# ================= INTERNAL HELPERS
def _get_table():
    doc = Document(WORD_FILE)
    REQUIRED = {"date", "status", "hard topic", "c topic", "java topic"}

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip().lower() for c in row.cells]
            if REQUIRED.issubset(set(cells)):
                headers = {
                    c.text.strip().lower(): i
                    for i, c in enumerate(row.cells)
                }
                return table, headers

    raise ValueError("Required table not found")

def _get_all_rows():
    doc = Document(WORD_FILE)
    all_rows = []

    REQUIRED = {"date", "status", "hard topic", "c topic", "java topic"}

    for table in doc.tables:
        headers = None
        header_idx = None

        for i, row in enumerate(table.rows):
            cells = [c.text.strip().lower() for c in row.cells if c.text.strip()]
            if REQUIRED.issubset(set(cells)):
                headers = {
                    c.text.strip().lower(): idx
                    for idx, c in enumerate(row.cells)
                    if c.text.strip()
                }
                header_idx = i
                break

        if not headers:
            continue

        for row in table.rows[header_idx + 1:]:
            # skip empty rows
            if not row.cells[headers["date"]].text.strip():
                continue
            all_rows.append((row, headers))

    if not all_rows:
        raise ValueError("No tables with required columns found.")

    return all_rows

def _parse_date(row, headers):
    text = row.cells[headers["date"]].text

    if not text:
        return None

    # Clean Word formatting junk
    text = text.strip().replace("\n", "").replace("\xa0", "")

    try:
        return datetime.strptime(
            text[:10], "%Y-%m-%d"
        ).date()
    except:
        return None


# ================= PDF CORE =================
def generate_pdf(start_date, end_date, title):
    all_rows = _get_all_rows()
    pdf_path = f"{title.replace(' ', '_')}.pdf"
    c = canvas.Canvas(pdf_path, pagesize=A4)
    width, height = A4

    # Title
    y = height - 2 * cm
    c.setFont("Helvetica-Bold", 14)
    c.drawString(2 * cm, y, title)
    y -= 1.5 * cm

    # Table header
    c.setFont("Helvetica-Bold", 10)
    c.setFillColor(colors.black)
    c.drawString(2 * cm, y, "Date")
    c.drawString(5 * cm, y, "Status")
    c.drawString(7 * cm, y, "Hard Topic")
    c.drawString(11 * cm, y, "C Topic")
    c.drawString(15 * cm, y, "Java Topic")
    y -= 1 * cm

    c.setFont("Helvetica", 10)

    for row, headers in all_rows:
        d = _parse_date(row, headers)
        if not d or not (start_date <= d <= end_date):
            continue

        raw_status = row.cells[headers["status"]].text.strip()
        status = "DONE" if "✅" in raw_status else "MISS" if "❌" in raw_status else "-"

        hard = row.cells[headers["hard topic"]].text.strip() or "None"
        c_topic = row.cells[headers["c topic"]].text.strip() or "-"
        java_topic = row.cells[headers["java topic"]].text.strip() or "-"

        # Column X positions
        x_date = 2 * cm
        x_status = 5 * cm
        x_hard = 7 * cm
        x_c = 11 * cm
        x_java = 15 * cm

        # Date
        c.setFillColor(colors.black)
        c.drawString(x_date, y, str(d))

        # Status (ONLY this colored)
        if status == "DONE":
            c.setFillColor(colors.green)
        elif status == "MISS":
            c.setFillColor(colors.red)
        else:
            c.setFillColor(colors.black)
        c.drawString(x_status, y, status)

        # Rest (black)
        c.setFillColor(colors.black)
        c.drawString(x_hard, y, hard[:25])
        c.drawString(x_c, y, c_topic[:25])
        c.drawString(x_java, y, java_topic[:25])

        y -= 0.8 * cm

        # Page break
        if y < 2 * cm:
            c.showPage()
            y = height - 2 * cm
            c.setFont("Helvetica-Bold", 10)
            c.drawString(2 * cm, y, "Date")
            c.drawString(5 * cm, y, "Status")
            c.drawString(7 * cm, y, "Hard Topic")
            c.drawString(11 * cm, y, "C Topic")
            c.drawString(15 * cm, y, "Java Topic")
            y -= 1 * cm
            c.setFont("Helvetica", 10)

    c.save()
    return pdf_path

# ================= WEEKLY PDF =================
async def send_weekly_report(context):
    today = datetime.now(IST).date()
    start = today - timedelta(days=6)
    pdf = generate_pdf(start, today, "Weekly Study Report")
    await context.bot.send_document(chat_id=CHAT_ID, document=open(pdf, "rb"))


# ================= MONTHLY PDF =================
async def send_monthly_report(context):
    today = datetime.now(IST).date()
    first = today.replace(day=1)
    last_end = first - timedelta(days=1)
    last_start = last_end.replace(day=1)

    pdf = generate_pdf(
        last_start,
        last_end,
        f"Monthly Report {last_start.strftime('%B %Y')}",
    )
    await context.bot.send_document(chat_id=CHAT_ID, document=open(pdf, "rb"))


async def monthly_checker(context):
    if (datetime.now(IST).date() + timedelta(days=1)).day == 1:
        await send_monthly_report(context)


# ================= SUNDAY SUMMARY =================
async def send_sunday_summary(context):
    today = datetime.now(IST).date()
    start = today - timedelta(days=6)

    table, headers = _get_table()
    done = miss = 0
    hard_topics = []

    for row in table.rows:
        d = _parse_date(row, headers)
        if not d or not (start <= d <= today):
            continue

        s = row.cells[headers["status"]].text
        h = row.cells[headers["hard topic"]].text

        if "✅" in s:
            done += 1
        elif "❌" in s:
            miss += 1

        if h and h.lower() != "none":
            hard_topics.append(f"• {d}: {h}")

    msg = (
        f"📌 *Weekly Summary*\n\n"
        f"✅ Done: {done}\n"
        f"❌ Missed: {miss}\n\n"
        f"🧠 *Hard Topics*\n"
        + ("\n".join(hard_topics) if hard_topics else "None 🎉")
    )

    await context.bot.send_message(chat_id=CHAT_ID, text=msg, parse_mode="Markdown")


# ================= GRAPHS =================
async def _plot_and_send(x, y, title, file, caption, context):
    plt.figure()
    plt.plot(x, y, marker="o")
    plt.yticks([0, 1], ["❌", "✅"])
    plt.title(title)
    plt.grid(True)
    plt.savefig(file)
    plt.close()

    await context.bot.send_photo(
    chat_id=CHAT_ID,
    photo=open(file, "rb"),
    caption=caption
    )

async def send_weekly_graph(context):
    today = datetime.now(IST).date()
    start = today - timedelta(days=6)

    table, headers = _get_table()
    x, y = [], []

    for row in table.rows:
        d = _parse_date(row, headers)
        if not d or not (start <= d <= today):
            continue

        x.append(d.strftime("%a"))
        y.append(1 if "✅" in row.cells[headers["status"]].text else 0)

    if x:
        await _plot_and_send(x, y, "Weekly Progress", "weekly.png", "📈 Weekly Progress", context)


async def send_monthly_graph(context):
    today = datetime.now(IST).date()
    first = today.replace(day=1)
    lm_end = first - timedelta(days=1)
    lm_start = lm_end.replace(day=1)

    table, headers = _get_table()
    x, y = [], []

    for row in table.rows:
        d = _parse_date(row, headers)
        if not d or not (lm_start <= d <= lm_end):
            continue

        x.append(d.day)
        y.append(1 if "✅" in row.cells[headers["status"]].text else 0)

    if x:
        await _plot_and_send(
            x,
            y,
            f"Monthly Progress {lm_start.strftime('%B')}",
            "monthly.png",
            "📊 Monthly Progress",
            context,
        )


# ================= CONSISTENCY =================
async def send_consistency_score(context):
    table, headers = _get_table()
    total = done = 0

    for row in table.rows:
        s = row.cells[headers["status"]].text
        if "✅" in s or "❌" in s:
            total += 1
            if "✅" in s:
                done += 1

    if not total:
        await context.bot.send_message(
            chat_id=CHAT_ID,
            text="📈 *Consistency*: No data yet",
            parse_mode="Markdown"
        )
        return
    pct = round((done / total) * 100, 2)
    await context.bot.send_message(
        chat_id=CHAT_ID, text=f"📈 *Consistency*: {pct}%", parse_mode="Markdown"
    )

# ================= BEST STREAK =================
async def send_best_streak(context):
    table, headers = _get_table()
    best = cur = 0

    for row in table.rows:
        if "✅" in row.cells[headers["status"]].text:
            cur += 1
            best = max(best, cur)
        else:
            cur = 0

    await context.bot.send_message(
        chat_id=CHAT_ID, text=f"🏆 *Best Streak*: {best} days", parse_mode="Markdown"
    )


# ================= STUDY SCORE =================
async def send_study_score(context):
    table, headers = _get_table()
    score = max_score = 0

    for row in table.rows:
        s = row.cells[headers["status"]].text
        h = row.cells[headers["hard topic"]].text.lower()

        if "✅" in s or "❌" in s:
            max_score += 10
            if "✅" in s:
                score += max(10 - (2 if h != "none" else 0), 0)

    if not max_score:
        await context.bot.send_message(
            chat_id=CHAT_ID,
            text="🧮 *Study Score*: No data yet",
            parse_mode="Markdown"
        )
        return
    pct = round((score / max_score) * 100, 2)
    await context.bot.send_message(
        chat_id=CHAT_ID, text=f"🧮 *Study Score*: {pct}%", parse_mode="Markdown"
    )       
# ================= HARD TOPIC ANALYTICS =================
async def send_hard_topic_analytics(context):
    table, headers = _get_table()
    topics = []

    for row in table.rows:
        h = row.cells[headers["hard topic"]].text
        if h and h.lower() != "none":
            topics.append(h)

    if not topics:
        msg = "🧠 *Hard Topics*: None 🎉"
    else:
        c = Counter(topics).most_common(5)
        msg = "🧠 *Hard Topic Analytics*\n\n" + "\n".join(
            f"• {t} → {n}" for t, n in c
        )

    await context.bot.send_message(chat_id=CHAT_ID, text=msg, parse_mode="Markdown")


# ================= MONTH COMPARISON =================
async def send_month_comparison(context):
    table, headers = _get_table()
    today = datetime.now(IST).date()

    first = today.replace(day=1)
    lm_end = first - timedelta(days=1)
    lm_start = lm_end.replace(day=1)
    pm_end = lm_start - timedelta(days=1)
    pm_start = pm_end.replace(day=1)

    def count(start, end):
        c = 0
        for row in table.rows:
            d = _parse_date(row, headers)
            if d and start <= d <= end and "✅" in row.cells[headers["status"]].text:
                c += 1
        return c

    lm, pm = count(lm_start, lm_end), count(pm_start, pm_end)
    trend = "📈 Improved" if lm > pm else "📉 Declined" if lm < pm else "➖ Same"

    msg = (
        f"🏅 *Month Comparison*\n\n"
        f"{pm_start.strftime('%B')}: {pm}\n"
        f"{lm_start.strftime('%B')}: {lm}\n\n"
        f"Trend: *{trend}*"
    )

    await context.bot.send_message(chat_id=CHAT_ID, text=msg, parse_mode="Markdown")


# ================= AI MOTIVATION =================
async def send_ai_motivation(context):
    table, headers = _get_table()
    score = max_score = 0

    for row in table.rows:
        s = row.cells[headers["status"]].text
        h = row.cells[headers["hard topic"]].text.lower()

        if "✅" in s or "❌" in s:
            max_score += 10
            if "✅" in s:
                score += max(10 - (2 if h != "none" else 0), 0)

    if not max_score:
        await context.bot.send_message(
            chat_id=CHAT_ID,
            text="🤖 *AI Motivation*\n\nNo study data yet. Let’s start strong 💪",
            parse_mode="Markdown"
        )
        return

    pct = (score / max_score) * 100

    if pct >= 85:
        msg = "🔥 Elite consistency! Keep dominating 🚀"
    elif pct >= 70:
        msg = "💪 Great job! Push a little harder."
    elif pct >= 50:
        msg = "🙂 Decent start. Reduce missed days."
    else:
        msg = "⚠️ Reset time. Small wins daily."

    await context.bot.send_message(
        chat_id=CHAT_ID, text=f"🤖 *AI Motivation*\n\n{msg}", parse_mode="Markdown"
    )


# ================= COMMAND =================
async def report_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    start = datetime.strptime(context.args[0], "%Y-%m-%d").date()
    end = datetime.strptime(context.args[1], "%Y-%m-%d").date()
    pdf = generate_pdf(start, end, f"Study Report {start} to {end}")
    await update.message.reply_document(open(pdf, "rb"))

# ================= TUESDAY MANUAL WEEKLY =================
async def manual_weekly_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    # Tuesday = 1
    if datetime.now(IST).weekday() != 6:
        await update.message.reply_text(
            "❌ Manual weekly reports are only available on Sunday."
        )
        return
    await send_weekly_report(context)
    await send_sunday_summary(context)
    await send_weekly_graph(context)
    await send_consistency_score(context)
    await send_best_streak(context)
    await send_study_score(context)
    await send_hard_topic_analytics(context)
    await send_ai_motivation(context)
    await update.message.reply_text("✅ Weekly reports sent manually.")
async def monday_bundle(context):
    await send_weekly_report(context)
    await send_sunday_summary(context)
    await send_weekly_graph(context)
    await send_consistency_score(context)
    await send_best_streak(context)
    await send_study_score(context)
    await send_hard_topic_analytics(context)
    await send_ai_motivation(context)
    await send_monthly_graph(context)
    await send_month_comparison(context)

# ================= REGISTER =================
def register_reports(app):
    app.add_handler(CommandHandler("report", report_command))
    app.add_handler(CommandHandler("weekly", manual_weekly_command))
    if app.job_queue is None:
        print("❌ JobQueue not available")
        return
    app.job_queue.run_daily(
    monday_bundle,
    dt_time(20, 52, tzinfo=IST), 
    days=(JobQueue.SUNDAY,), 
    )